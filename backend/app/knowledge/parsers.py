from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import fitz
from docx import Document


@dataclass(frozen=True)
class NormalizedAsset:
    local_id: str
    asset_type: str
    page_number: int
    content: bytes
    extension: str
    bbox: tuple[float, float, float, float] | None = None
    caption: str = ""
    ocr_text: str = ""


@dataclass(frozen=True)
class NormalizedBlock:
    kind: str
    text: str
    heading_level: int | None = None
    bbox: tuple[float, float, float, float] | None = None
    asset_local_id: str | None = None


@dataclass(frozen=True)
class NormalizedPage:
    page_number: int
    blocks: list[NormalizedBlock] = field(default_factory=list)


@dataclass(frozen=True)
class NormalizedDocument:
    title: str
    pages: list[NormalizedPage]
    assets: list[NormalizedAsset] = field(default_factory=list)


class ParserRegistry:
    def parse(self, path: Path, mime_type: str) -> NormalizedDocument:
        if mime_type in {"text/plain", "text/markdown"}:
            return self._parse_text(path, markdown=mime_type == "text/markdown" or path.suffix.lower() == ".md")
        if mime_type == "application/pdf":
            return self._parse_pdf(path)
        if mime_type.endswith("wordprocessingml.document"):
            return self._parse_docx(path)
        if mime_type in {"image/png", "image/jpeg"}:
            return self._parse_image(path, mime_type)
        raise ValueError(f"unsupported parser mime: {mime_type}")

    @staticmethod
    def _parse_text(path: Path, *, markdown: bool) -> NormalizedDocument:
        text = path.read_text(encoding="utf-8")
        blocks: list[NormalizedBlock] = []
        paragraph: list[str] = []

        def flush_paragraph() -> None:
            if paragraph:
                blocks.append(NormalizedBlock(kind="paragraph", text="\n".join(paragraph).strip()))
                paragraph.clear()

        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                flush_paragraph()
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if markdown and heading:
                flush_paragraph()
                blocks.append(NormalizedBlock(kind="heading", text=heading.group(2).strip(), heading_level=len(heading.group(1))))
            elif line.startswith(("- ", "* ")):
                flush_paragraph()
                blocks.append(NormalizedBlock(kind="list", text=line[2:].strip()))
            else:
                paragraph.append(line)
        flush_paragraph()
        return NormalizedDocument(title=path.stem, pages=[NormalizedPage(page_number=1, blocks=blocks)])

    @staticmethod
    def _parse_pdf(path: Path) -> NormalizedDocument:
        pages: list[NormalizedPage] = []
        assets: list[NormalizedAsset] = []
        with fitz.open(path) as document:
            for page_index, page in enumerate(document, start=1):
                blocks: list[NormalizedBlock] = []
                for item in page.get_text("blocks"):
                    text = str(item[4]).strip()
                    if not text:
                        continue
                    kind = "heading" if len(text) <= 80 and text.count("\n") == 0 and text[-1:] not in "。.!?；;" else "paragraph"
                    blocks.append(
                        NormalizedBlock(
                            kind=kind,
                            text=text,
                            heading_level=1 if kind == "heading" else None,
                            bbox=(float(item[0]), float(item[1]), float(item[2]), float(item[3])),
                        )
                    )
                for image_index, image in enumerate(page.get_images(full=True), start=1):
                    extracted = document.extract_image(image[0])
                    local_id = f"page-{page_index}-image-{image_index}"
                    assets.append(
                        NormalizedAsset(
                            local_id=local_id,
                            asset_type="image",
                            page_number=page_index,
                            content=extracted["image"],
                            extension=f".{extracted.get('ext', 'png')}",
                        )
                    )
                    blocks.append(NormalizedBlock(kind="image", text=f"图片 {image_index}", asset_local_id=local_id))
                pages.append(NormalizedPage(page_number=page_index, blocks=blocks))
        return NormalizedDocument(title=path.stem, pages=pages, assets=assets)

    @staticmethod
    def _parse_docx(path: Path) -> NormalizedDocument:
        document = Document(path)
        blocks: list[NormalizedBlock] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower()
            heading_match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style)
            if heading_match:
                level = int(next(group for group in heading_match.groups() if group))
                blocks.append(NormalizedBlock(kind="heading", text=text, heading_level=level))
            elif "list" in style or "列表" in style:
                blocks.append(NormalizedBlock(kind="list", text=text))
            else:
                blocks.append(NormalizedBlock(kind="paragraph", text=text))
        for table in document.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            if rows:
                blocks.append(NormalizedBlock(kind="table", text="\n".join(rows)))
        return NormalizedDocument(title=path.stem, pages=[NormalizedPage(page_number=1, blocks=blocks)])

    @staticmethod
    def _parse_image(path: Path, mime_type: str) -> NormalizedDocument:
        extension = ".png" if mime_type == "image/png" else ".jpg"
        asset = NormalizedAsset(
            local_id="image-1",
            asset_type="image",
            page_number=1,
            content=path.read_bytes(),
            extension=extension,
            caption=path.stem,
        )
        block = NormalizedBlock(kind="image", text=path.stem, asset_local_id=asset.local_id)
        return NormalizedDocument(title=path.stem, pages=[NormalizedPage(page_number=1, blocks=[block])], assets=[asset])
